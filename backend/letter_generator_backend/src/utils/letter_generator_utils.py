import pandas as pd
import re
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def clean_contractor_name(name):
    """Удаляет первые 10 цифр из названия контрагента"""
    if isinstance(name, str):
        # Удаляем первые 10 значное число в начале
        cleaned = re.sub(r'^\d{10}\s*', '', name).strip()
        return cleaned
    return name

def get_contractor_short_name(full_name):
    """Получает сокращенное наименование контрагента"""
    if not full_name:
        return ""
    
    # Ищем ООО, АО, ЗАО и т.д.
    patterns = [
        r'Общество с ограниченной ответственностью\s*"([^"]+)"',
        r'ООО\s*"([^"]+)"',
        r'Акционерное общество\s*"([^"]+)"',
        r'АО\s*"([^"]+)"',
        r'Закрытое акционерное общество\s*"([^"]+)"',
        r'ЗАО\s*"([^"]+)"'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, full_name, re.IGNORECASE)
        if match:
            return match.group(1)
    
    # Если не найдено, возвращаем исходное название
    return full_name

def get_contractor_full_form(name):
    """Определяет полную форму организации"""
    if not name:
        return ""
    
    name_lower = name.lower()
    if 'ооо' in name_lower or 'общество с ограниченной ответственностью' in name_lower:
        return 'Обществом с ограниченной ответственностью'
    elif 'ао' in name_lower or 'акционерное общество' in name_lower:
        return 'Акционерным обществом'
    elif 'зао' in name_lower or 'закрытое акционерное общество' in name_lower:
        return 'Закрытым акционерным обществом'
    else:
        return 'организацией'

def calculate_penalty(amount, days_overdue):
    """Расчет пени с учетом сложного процента"""
    if days_overdue <= 0:
        return 0
    
    penalty = 0
    remaining_days = days_overdue
    current_amount = float(amount)
    
    # Первые 14 дней - 0.1% в день
    if remaining_days > 0:
        days_first_period = min(remaining_days, 14)
        for day in range(days_first_period):
            daily_penalty = current_amount * 0.001  # 0.1%
            penalty += daily_penalty
            current_amount += daily_penalty
        remaining_days -= days_first_period
    
    # После 14 дней - 0.5% в день
    if remaining_days > 0:
        for day in range(remaining_days):
            daily_penalty = current_amount * 0.005  # 0.5%
            penalty += daily_penalty
            current_amount += daily_penalty
    
    return penalty

def number_to_words_russian(number):
    """Преобразование числа в слова на русском языке"""
    
    def num_to_words(n):
        if n == 0:
            return "ноль"
        
        ones = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять",
                "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать",
                "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
        
        tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
        
        hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
        
        def convert_hundreds(num):
            result = []
            if num >= 100:
                result.append(hundreds[num // 100])
                num %= 100
            if num >= 20:
                result.append(tens[num // 10])
                num %= 10
            if num > 0:
                result.append(ones[num])
            return " ".join(result)
        
        if n < 1000:
            return convert_hundreds(n)
        elif n < 1000000:
            thousands_part = n // 1000
            remainder = n % 1000
            result = convert_hundreds(thousands_part)
            if thousands_part % 10 == 1 and thousands_part % 100 != 11:
                result += " тысяча"
            elif thousands_part % 10 in [2, 3, 4] and thousands_part % 100 not in [12, 13, 14]:
                result += " тысячи"
            else:
                result += " тысяч"
            if remainder > 0:
                result += " " + convert_hundreds(remainder)
            return result
        else:
            millions_part = n // 1000000
            remainder = n % 1000000
            result = convert_hundreds(millions_part)
            if millions_part % 10 == 1 and millions_part % 100 != 11:
                result += " миллион"
            elif millions_part % 10 in [2, 3, 4] and millions_part % 100 not in [12, 13, 14]:
                result += " миллиона"
            else:
                result += " миллионов"
            if remainder > 0:
                if remainder >= 1000:
                    thousands_part = remainder // 1000
                    hundreds_part = remainder % 1000
                    result += " " + convert_hundreds(thousands_part)
                    if thousands_part % 10 == 1 and thousands_part % 100 != 11:
                        result += " тысяча"
                    elif thousands_part % 10 in [2, 3, 4] and thousands_part % 100 not in [12, 13, 14]:
                        result += " тысячи"
                    else:
                        result += " тысяч"
                    if hundreds_part > 0:
                        result += " " + convert_hundreds(hundreds_part)
                else:
                    result += " " + convert_hundreds(remainder)
            return result
    
    return num_to_words(int(number))

def format_amount_in_words(amount):
    """Форматирование суммы прописью"""
    rubles = int(amount)
    kopecks = int(round((amount - rubles) * 100))
    
    rubles_words = number_to_words_russian(rubles)
    
    # Склонение рублей
    if rubles % 10 == 1 and rubles % 100 != 11:
        ruble_form = "рубль"
    elif rubles % 10 in [2, 3, 4] and rubles % 100 not in [12, 13, 14]:
        ruble_form = "рубля"
    else:
        ruble_form = "рублей"
    
    # Склонение копеек
    if kopecks % 10 == 1 and kopecks % 100 != 11:
        kopeck_form = "копейка"
    elif kopecks % 10 in [2, 3, 4] and kopecks % 100 not in [12, 13, 14]:
        kopeck_form = "копейки"
    else:
        kopeck_form = "копеек"
    
    return f"{rubles_words} {ruble_form} {kopecks:02d} {kopeck_form}"

def process_reporting_data(reporting_path, sed_path):
    """Обработка данных из файлов отчетности и СЭД"""
    try:
        # Читаем файл отчетности
        reporting_df = pd.read_excel(reporting_path, sheet_name='Внутригрупповая отчетность')
        
        # Читаем файл СЭД
        sed_df = pd.read_excel(sed_path)
        
        current_date = datetime.now()
        processed_data = {}
        
        # Обрабатываем каждую строку отчетности
        for index, row in reporting_df.iterrows():
            try:
                order_number = row.iloc[6] if len(row) > 6 else None  # Колонка G (индекс 6)
                contractor_name = row.iloc[9] if len(row) > 9 else None  # Колонка J (индекс 9)
                planned_date = row.iloc[17] if len(row) > 17 else None  # Колонка R (индекс 17)
                actual_date = row.iloc[29] if len(row) > 29 else None  # Колонка AD (индекс 29)
                amount = row.iloc[16] if len(row) > 16 else 0  # Колонка Q (индекс 16)
                
                # Дополнительные поля для приложения
                col_k = row.iloc[10] if len(row) > 10 else ""  # Колонка K
                col_l = row.iloc[11] if len(row) > 11 else ""  # Колонка L
                col_m = row.iloc[12] if len(row) > 12 else ""  # Колонка M
                col_n = row.iloc[13] if len(row) > 13 else ""  # Колонка N
                col_p = row.iloc[15] if len(row) > 15 else ""  # Колонка P
                
                if pd.isna(order_number) or pd.isna(planned_date) or pd.isna(contractor_name):
                    continue
                
                # Преобразуем даты в datetime объекты, обрабатывая NaT
                planned_date_dt = pd.to_datetime(planned_date)
                actual_date_dt = pd.to_datetime(actual_date) if pd.notna(actual_date) else None

                if pd.isna(planned_date_dt):
                    continue # Пропускаем строки без валидной плановой даты

                # Определяем просрочку
                is_overdue = False
                days_overdue = 0
                category = ''
                
                if actual_date_dt is None:
                    # Нет фактической даты - сравниваем с текущей датой
                    if current_date > planned_date_dt:
                        is_overdue = True
                        days_overdue = (current_date - planned_date_dt).days
                        category = 'просрочено не поставлено'
                else:
                    # Есть фактическая дата
                    if actual_date_dt > planned_date_dt:
                        is_overdue = True
                        days_overdue = (actual_date_dt - planned_date_dt).days
                        category = 'поставленные просрочки'
                
                if is_overdue and amount > 0:
                    # Очищаем название контрагента
                    clean_contractor = clean_contractor_name(contractor_name)
                    
                    # Ищем данные в файле СЭД
                    sed_row = sed_df[sed_df.iloc[:, 5] == order_number]  # Колонка F (индекс 5)
                    
                    if not sed_row.empty:
                        be_name = sed_row.iloc[0, 2] if len(sed_row.iloc[0]) > 2 else ""  # Колонка C
                        reg_number = str(sed_row.iloc[0, 7]) if len(sed_row.iloc[0]) > 7 and pd.notna(sed_row.iloc[0, 7]) else ""  # Колонка H
                        reg_date_raw = sed_row.iloc[0, 15] if len(sed_row.iloc[0]) > 15 else None  # Колонка P
                        reg_date = pd.to_datetime(reg_date_raw).strftime('%d.%m.%Y') if pd.notna(reg_date_raw) else ""
                        
                        # Группируем по контрагенту и заказу
                        key = f"{clean_contractor}_{order_number}"
                        
                        if key not in processed_data:
                            processed_data[key] = {
                                'order_number': order_number,
                                'contractor_name': clean_contractor,
                                'contractor_short_name': get_contractor_short_name(clean_contractor),
                                'contractor_full_form': get_contractor_full_form(clean_contractor),
                                'be_name': be_name,
                                'reg_number': reg_number,
                                'reg_date': reg_date,
                                'planned_date': planned_date_dt.strftime('%d.%m.%Y'),
                                'total_amount': 0,
                                'total_penalty': 0,
                                'total_positions': 0,
                                'category': category,
                                'positions': []
                            }
                        
                        # Рассчитываем пени для позиции
                        penalty = calculate_penalty(amount, days_overdue)
                        
                        # Добавляем позицию
                        position_data = {
                            'col_k': col_k,
                            'col_l': col_l,
                            'col_m': col_m,
                            'col_n': col_n,
                            'col_p': col_p,
                            'amount': amount,
                            'days_overdue': days_overdue,
                            'penalty': penalty
                        }
                        
                        processed_data[key]['positions'].append(position_data)
                        processed_data[key]['total_amount'] += amount
                        processed_data[key]['total_penalty'] += penalty
                        processed_data[key]['total_positions'] += 1
                        
            except Exception as e:
                print(f"Ошибка обработки строки {index}: {str(e)}")
                continue
        
        return list(processed_data.values())
        
    except Exception as e:
        raise Exception(f"Ошибка при обработке файлов: {str(e)}")

def generate_letter_document(letter_data, output_path):
    """Генерация документа письма"""
    try:
        doc = Document()
        
        # Добавляем логотип в верхний колонтитул
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        # Предполагаем, что логотип находится в папке static
        # В реальном проекте нужно будет указать путь к логотипу
        try:
            run.add_picture(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'static', 'logo.png'), width=Inches(1.0))
        except Exception as e:
            print(f"Не удалось добавить логотип: {e}")
        
        # Заголовок (Номер и Кас) - слева
        header_paragraph = doc.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.add_run(f"№ ____________\nКас.: Претензионная работа по договору поставки")
        
        # Обращение - по центру и жирным
        salutation = doc.add_paragraph()
        salutation.alignment = WD_ALIGN_PARAGRAPH.CENTER
        salutation_run = salutation.add_run("Уважаемый партнер!")
        salutation_run.bold = True
        
        # Основной текст письма с жирным выделением сумм и количества
        main_text_parts = [
            f"Настоящим сообщаем, что между «{letter_data['be_name']}» и {letter_data['contractor_full_form']} «{letter_data['contractor_name']}» (далее – «{letter_data['contractor_short_name']}») заключен договор поставки № {letter_data['reg_number']} от {letter_data['reg_date']} (далее – Договор поставки). В соответствии с Договором поставки сторонами подписана Спецификация № ",
            letter_data['order_number'],
            f" от 03.03.2025 (далее – спецификация), согласно которой «{letter_data['contractor_short_name']}» обязуется в срок до {letter_data['planned_date']} поставить товары на сумму ",
            f"{letter_data['total_amount']:.2f} ({format_amount_in_words(letter_data['total_amount'])})",
            f", а «{letter_data['be_name']}» - оплатить указанные товары в течение 30 (тридцати) календарных дней с момента их передачи (Приложение № 1 к настоящему письму).\n\nПо состоянию на {datetime.now().strftime('%d.%m.%Y')} товары в количестве ",
            str(letter_data['total_positions']),
            f" позиций на ",
            f"{letter_data['total_amount']:.2f} ({format_amount_in_words(letter_data['total_amount'])})",
            f" в месте поставки {'отсутствуют' if letter_data['category'] == 'просрочено не поставлено' else 'поступили с просрочкой'}, что является нарушением п. 4.1 Договора поставки. На основании п. 8.3. Договора поставки сумма пени на текущий момент по просроченным позициям составляет ",
            f"{letter_data['total_penalty']:.2f} ({format_amount_in_words(letter_data['total_penalty'])})",
            " и рассчитывается следующим образом:\n\n0,1 (Ноль целых и одна десятая) % стоимости непоставленного в срок товара, или товара, в отношении которого не выполнены требования, предъявленные Покупателем в соответствии с пунктами 7.5. и 7.10.5. договора, за каждый день просрочки в течение первых двух недель, а в случае дальнейшей просрочки - в размере 0,5 (Ноль целых и пять десятых) % стоимости такого товара за каждый день просрочки.\n\n",
            f"Обращаем Ваше внимание на то, что в настоящее время имеется перечень критичных для «{letter_data['be_name']}» позиций товара (Приложение № 2 к настоящему письму), поставка которых должна быть осуществлена до {letter_data['planned_date']}, при этом, риски срыва сроков поставки являются недопустимыми.\n\nУчитывая изложенное, убедительно просим Вас ускорить исполнение обязательств, принятых по Договору поставки, в части своевременной отгрузки товаров и поставки товара в целях недопущения увеличения суммы пени по позициям товара согласно Приложению № 1 к настоящему письму и минимизации рисков образования пени по позициям товаров согласно приложению № 2 к письму."
        ]
        
        # Добавляем основной текст с жирным выделением
        main_paragraph = doc.add_paragraph()
        bold_indices = [1, 3, 5, 7, 9]  # Индексы элементов, которые нужно выделить жирным
        
        for i, part in enumerate(main_text_parts):
            run = main_paragraph.add_run(part)
            if i in bold_indices:
                run.bold = True
        
        # Приложения
        doc.add_paragraph("\nПриложения по тексту:")
        doc.add_paragraph(f"1) Спецификация № {letter_data['reg_number']} от {letter_data['reg_date']} (на 6 л. в 1 экз.);")
        doc.add_paragraph(f"2) Спецификация № {letter_data['order_number']} от {datetime.now().strftime('%d.%m.%Y')} (на {len(letter_data['positions'])} л. в 1 экз.)")
        
        # Подпись
        signature = doc.add_paragraph("\n\nС уважением,")
        signature.add_run("\n\n[_____________________] [_____________]")
        signature.add_run("\n[_________________________] (подпись) (Ф.И.О. уполномоченного")
        signature.add_run("\n(наименование должности уполномоченного м.п. лица УК/УО на подписание")
        signature.add_run("\nлица УК/УО] )")
        
        # Исполнитель
        doc.add_paragraph("\n\nИсп. [______________________________________]")
        doc.add_paragraph("(Ф.И.О. Отв. Исполнителя УК/УО)")
        doc.add_paragraph("Контактный т.[_______________________________]")
        doc.add_paragraph("(контактный номер телефона Отв. Исполнителя УК/УО)")
        
        # Сохраняем документ
        doc.save(output_path)
        
        return True
        
    except Exception as e:
        raise Exception(f"Ошибка при генерации документа: {str(e)}")

def generate_appendix_document(letter_data, output_path):
    """Генерация приложения к письму"""
    try:
        doc = Document()
        
        # Заголовок приложения
        doc.add_heading(f'Приложение № 1 к письму', 0)
        doc.add_heading(f'Спецификация по заказу № {letter_data["order_number"]}', 1)
        
        # Шапка с жирным выделением
        header_info = doc.add_paragraph()
        header_info.add_run(f"Номер заказа: ")
        run_order = header_info.add_run(letter_data['order_number'])
        run_order.bold = True
        header_info.add_run(f"\nКоличество просроченных позиций: ")
        run_positions = header_info.add_run(str(letter_data['total_positions']))
        run_positions.bold = True
        header_info.add_run(f"\nНа сумму: ")
        run_amount = header_info.add_run(f"{letter_data['total_amount']:.2f} ({format_amount_in_words(letter_data['total_amount'])})")
        run_amount.bold = True
        
        # Таблица с позициями
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Колонка K'
        hdr_cells[1].text = 'Колонка L'
        hdr_cells[2].text = 'Колонка M'
        hdr_cells[3].text = 'Колонка N'
        hdr_cells[4].text = 'Колонка Q (Сумма)'
        hdr_cells[5].text = 'Колонка P'
        hdr_cells[6].text = 'Дни просрочки'
        
        # Добавляем строки с данными
        for position in letter_data['positions']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(position['col_k'])
            row_cells[1].text = str(position['col_l'])
            row_cells[2].text = str(position['col_m'])
            row_cells[3].text = str(position['col_n'])
            row_cells[4].text = f"{position['amount']:.2f}"
            row_cells[5].text = str(position['col_p'])
            row_cells[6].text = str(position['days_overdue'])
        
        # Сохраняем документ
        doc.save(output_path)
        
        return True
        
    except Exception as e:
        raise Exception(f"Ошибка при генерации приложения: {str(e)}")

